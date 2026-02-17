# app.py
import re
import time
from dataclasses import dataclass
from typing import Optional, List, Tuple, Dict, Iterable
from collections import Counter, defaultdict
from io import BytesIO
from datetime import datetime
import json
import os
import tempfile

import streamlit as st
import pandas as pd
import requests
from rapidfuzz import fuzz, process
from tenacity import retry, stop_after_attempt, wait_exponential

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm

# For Word export (reformatted references)
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

# Optional readers (DOCX/PDF)
try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None


# =============================
# Usage counter system (Streamlit Cloud safe)
# =============================
USAGE_FILE = os.path.join(tempfile.gettempdir(), "usage_stats.json")


def load_usage_stats():
    default = {
        "app_runs": 0,
        "files_processed": 0,
        "citations_checked": 0,
        "references_checked": 0,
    }
    try:
        if not os.path.exists(USAGE_FILE):
            return default
        with open(USAGE_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        for k, v in default.items():
            data.setdefault(k, v)
        return data
    except Exception:
        return default


def save_usage_stats(stats):
    try:
        with open(USAGE_FILE, "w", encoding="utf-8") as f:
            json.dump(stats, f, indent=2)
    except Exception:
        pass


def increment_usage(app_run=False, file=False, cites=0, refs=0):
    stats = load_usage_stats()
    if app_run:
        stats["app_runs"] += 1
    if file:
        stats["files_processed"] += 1
    stats["citations_checked"] += int(cites or 0)
    stats["references_checked"] += int(refs or 0)
    save_usage_stats(stats)
    return stats


def increment_once_per_session(key: str, **kwargs):
    if st.session_state.get(key):
        return load_usage_stats()
    st.session_state[key] = True
    return increment_usage(**kwargs)


# =============================
# Constants
# =============================
YEAR = r"(?:19|20)\d{2}[a-z]?"  # allows 2020a
CROSSREF_API = "https://api.crossref.org/works"
OPENALEX_API = "https://api.openalex.org/works"

# include "Reference" (singular) because many manuscripts use it
REF_HEADINGS = [
    r"^\s*reference\s*[:|]?\s*$",
    r"^\s*references?\s*(?:list)?\s*[:|]?\s*$",
    r"^\s*bibliograph(?:y|ies)\s*[:|]?\s*$",
    r"^\s*works\s+cited\s*[:|]?\s*$",
    r"^\s*literature\s+cited\s*[:|]?\s*$",
]

# Expand org aliases
ORG_ALIASES: Dict[str, List[str]] = {
    "unctad": ["unctad", "united nations conference on trade and development"],
    "who": ["who", "world health organization", "world health organisation"],
    "oecd": [
        "oecd",
        "organisation for economic co-operation and development",
        "organization for economic cooperation and development",
    ],
    "world bank": ["world bank", "international bank for reconstruction and development", "ibrd"],
    "world bank group": ["world bank group"],
    "imf": ["imf", "international monetary fund"],
    "un": ["un", "united nations", "u.n.", "united nations organisation", "united nations organization"],
    "au": ["au", "african union"],
    "eu": ["eu", "european union"],
    "unesco": [
        "unesco",
        "united nations educational, scientific and cultural organization",
        "united nations educational scientific and cultural organization",
    ],
    "unicef": ["unicef", "united nations children's fund", "united nations childrens fund"],
}

ORG_ACRONYMS = {
    "UNCTAD",
    "WHO",
    "OECD",
    "IMF",
    "UN",
    "AU",
    "EU",
    "UNESCO",
    "UNICEF",
    "WORLD BANK",
    "WORLD BANK GROUP",
    "IBRD",
}

NONNAME_STOPWORDS = {
    "the", "a", "an", "and", "or", "of", "to", "in", "on", "for", "with", "as", "by",
    "from", "into", "at", "than", "that", "this", "these", "those",
    "recent", "minimum", "maximum", "mathematical", "basis", "adopting", "adopt", "model",
    "framework", "table", "figure", "chapter", "section", "appendix", "equation",
    "definition", "method", "methods", "result", "results", "discussion",
    "similarly", "moreover", "however", "therefore", "thus", "also", "further", "additionally",
    "consequently", "notably", "specifically", "generally", "overall", "first", "second", "finally",
    "public", "construct", "africa", "europe", "america", "world", "bank",
}

STYLE_APA = "APA-like"
STYLE_HARVARD = "Harvard-like"
STYLE_NUMERIC = "Numeric"


# =============================
# Data classes
# =============================
@dataclass
class InTextCitation:
    style: str
    raw: str
    key: str
    pretty: str
    author_or_org: Optional[str] = None
    year: Optional[str] = None
    number: Optional[int] = None


@dataclass
class ReferenceEntry:
    raw: str
    key: str
    pretty: str
    author_or_org: Optional[str] = None
    year: Optional[str] = None
    number: Optional[int] = None


# =============================
# Normalisation helpers
# =============================
def norm_spaces(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()


def norm_token(s: str) -> str:
    s = (s or "").lower().replace("’", "'")
    s = re.sub(r"[^a-z0-9\- ']", "", s)
    return norm_spaces(s)


def titleish(s: str) -> str:
    if not s:
        return s
    if s.strip().upper() in ORG_ACRONYMS:
        return s.strip().upper()
    words = s.strip().split()
    return " ".join((w[:1].upper() + w[1:]) if w else w for w in words)


def canon_org(name: str) -> str:
    n = norm_token(name)
    for canon, variants in ORG_ALIASES.items():
        for v in variants:
            if n == norm_token(v):
                return canon
    return n


def key_author_year(author_surname: str, year: str) -> str:
    return f"au_{norm_token(author_surname)}_{(year or '').lower()}"


def key_org_year(org: str, year: str) -> str:
    return f"org_{canon_org(org)}_{(year or '').lower()}"


def key_numeric(n: int) -> str:
    return f"n_{int(n)}"


def is_known_org(text: str) -> bool:
    t = (text or "").strip()
    if t.upper() in ORG_ACRONYMS:
        return True
    c = canon_org(t)
    return c in ORG_ALIASES.keys()


def looks_like_two_authors(name: str) -> bool:
    parts = (name or "").strip().split()
    return len(parts) == 2 and all(p[:1].isupper() for p in parts)


def looks_like_person_surname(token: str) -> bool:
    if not token:
        return False
    t = token.strip()

    if re.search(r"(?:'s|’s)$", t, flags=re.IGNORECASE):
        return False

    if not re.fullmatch(r"[A-Z][A-Za-z\-']{1,40}", t):
        return False

    if norm_token(t) in NONNAME_STOPWORDS:
        return False

    return True


def _first_surname_from_author_blob(blob: str) -> Optional[str]:
    if not blob:
        return None
    b = blob.strip()
    b = re.sub(r"^(see|e\.g\.|cf\.)\s+", "", b, flags=re.IGNORECASE).strip()
    b = re.sub(r"\s+et\s+al\.?$", "", b, flags=re.IGNORECASE).strip()

    first = b.split(",")[0].strip()
    first = re.split(r"\s+(?:and|&)\s+", first, flags=re.IGNORECASE)[0].strip()

    m = re.search(r"([A-Z][A-Za-z\-']+)", first)
    return m.group(1) if m else None


def format_author_blob_for_display(authors_blob: str) -> str:
    if not authors_blob:
        return ""
    s = authors_blob.strip()
    s = re.sub(r"\s*&\s*", " & ", s)
    s = re.sub(r"\s+and\s+", " and ", s, flags=re.IGNORECASE)
    s = re.sub(r"\s*,\s*", ", ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


# =============================
# DOCX block iterator
# =============================
def iter_block_items(document) -> Iterable[Tuple[str, str]]:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent = document.element.body
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            p = Paragraph(child, document)
            t = p.text.strip()
            if t:
                yield ("p", t)
        elif isinstance(child, CT_Tbl):
            tbl = Table(child, document)
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    yield ("t", " ".join(cells))


# =============================
# File readers
# =============================
def read_docx(file) -> str:
    if docx is None:
        raise RuntimeError("python-docx not installed")
    d = docx.Document(file)
    parts = []
    for _, txt in iter_block_items(d):
        parts.append(txt)
    # Keep paragraph boundaries
    return "\n".join(parts)


def read_pdf(file) -> str:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber not installed")
    out = []
    with pdfplumber.open(file) as pdf:
        for p in pdf.pages:
            out.append(p.extract_text() or "")
    return "\n".join(out)


# =============================
# Reference section detection
# =============================
def ref_line_score(line: str) -> int:
    s = (line or "").strip()
    if not s:
        return 0

    # numeric reference list: [1] ... or 1. ... or 1) ...
    if re.match(r"^\s*(\[\d+\]|\d+[\.\)])\s+", s):
        return 5

    # author-year list shapes
    if re.match(rf"^[A-Z][A-Za-z\-']+\s*,.*\(\s*{YEAR}\s*\)", s):
        return 6
    if re.match(rf"^[A-Z][A-Za-z\-']+\s*,.*\b{YEAR}\b", s):
        return 4

    # org-year list shapes
    if re.match(rf"^[A-Z][A-Z&\- ]{{2,}}\.\s*\(\s*{YEAR}\s*\)", s):
        return 6
    if re.match(rf"^[A-Z][A-Za-z&.\- ]{{2,}}?\s*\(\s*{YEAR}\s*\)", s):
        return 3

    return 0


def auto_detect_references_start(text: str) -> Tuple[int, float]:
    lines = (text or "").splitlines()
    if len(lines) < 40:
        idx = max(0, int(len(lines) * 0.70))
        return idx, 0.2

    window = 40
    start = int(len(lines) * 0.35)
    scores: List[Tuple[int, int]] = []
    for i in range(start, max(start + 1, len(lines) - window)):
        score = sum(ref_line_score(lines[j]) for j in range(i, i + window))
        scores.append((i, score))

    best_idx, best_score = max(scores, key=lambda x: x[1])
    confidence = min(1.0, best_score / window)
    return best_idx, confidence


def split_by_heading_or_autodetect(text: str) -> Tuple[str, str, str]:
    lines = (text or "").splitlines()

    heading_idx = None
    heading_line = ""
    for i, line in enumerate(lines):
        s = (line or "").strip()
        for pat in REF_HEADINGS:
            if re.search(pat, s, flags=re.IGNORECASE):
                heading_idx = i
                heading_line = s
                break
        if heading_idx is not None:
            break

    auto_idx, auto_conf = auto_detect_references_start(text)

    if heading_idx is None:
        chosen = auto_idx
        reason = f"No heading found, used auto-detect (confidence {auto_conf:.2f})."
    else:
        chosen = heading_idx + 1
        reason = f"Found heading: {heading_line}"

        # prefer later auto start if it looks safer (PDF spillover)
        if auto_conf >= 0.55 and auto_idx > heading_idx + 3:
            chosen = auto_idx
            reason = (
                f"Found heading: {heading_line}, but used auto-detect later start "
                f"(confidence {auto_conf:.2f}) to avoid spillover."
            )

    main = "\n".join(lines[:chosen]).strip()
    refs = "\n".join(lines[chosen:]).strip()
    return main, refs, reason


# =============================
# Reference splitting and parsing
# =============================
def _strip_leading_numbering(s: str) -> str:
    x = (s or "").strip()
    # [12] ...
    x = re.sub(r"^\s*\[\s*\d+\s*\]\s*", "", x)
    # 12. ... or 12) ...
    x = re.sub(r"^\s*\d+\s*[\.\)]\s*", "", x)
    return x.strip()


def split_reference_entries(ref_text: str) -> List[str]:
    """
    Robust splitter for reference lists.

    Works for:
    - numbered: [1] ... / 1. ... / 1) ...
    - unnumbered: one reference per paragraph/line (common in DOCX exports)
    - lightly wrapped lines (PDF extraction)
    """
    t = (ref_text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [ln.strip() for ln in t.split("\n") if ln.strip()]

    if not lines:
        return []

    # 1) numbered list
    start_pat_num = re.compile(r"^\s*(?:\[\s*\d+\s*\]|\d+\s*[\.\)])\s+")
    if any(start_pat_num.match(ln) for ln in lines[: min(25, len(lines))]):
        entries: List[str] = []
        buf: List[str] = []

        def flush():
            nonlocal buf
            if buf:
                e = norm_spaces(" ".join(buf)).strip()
                if e:
                    entries.append(e)
            buf = []

        for ln in lines:
            if start_pat_num.match(ln):
                flush()
                buf.append(ln)
            else:
                buf.append(ln)
        flush()
        return entries

    # 2) unnumbered but clearly "one reference per line"
    # If most lines look like reference-start lines, treat each line as an entry.
    start_pat_author = re.compile(rf"^[A-Z][A-Za-z\-']+,\s*.+?\(\s*{YEAR}\s*\)")
    start_pat_author2 = re.compile(rf"^[A-Z][A-Za-z\-']+,\s*.+?\b{YEAR}\b")
    start_pat_org = re.compile(rf"^[A-Z][A-Za-z&.\- ]{{2,}}?\s*\(\s*{YEAR}\s*\)")

    refish = 0
    for ln in lines[: min(80, len(lines))]:
        if start_pat_author.search(ln) or start_pat_author2.search(ln) or start_pat_org.search(ln):
            refish += 1
    if refish >= max(3, int(0.50 * min(80, len(lines)))):
        return [norm_spaces(ln) for ln in lines]

    # 3) fallback: detect starts and join wrapped lines until next start
    start_pat_loose = re.compile(
        rf"^(?:"
        rf"[A-Z][A-Za-z\-']+,\s*[A-Z]"      # surname, initial
        rf"|[A-Z][A-Za-z&.\- ]{{2,}}\s*\("  # org (year)
        rf")"
    )

    entries: List[str] = []
    buf: List[str] = []

    def flush2():
        nonlocal buf
        if buf:
            e = norm_spaces(" ".join(buf)).strip()
            if e:
                entries.append(e)
        buf = []

    for ln in lines:
        if start_pat_loose.match(ln) and buf:
            flush2()
        buf.append(ln)
    flush2()

    return entries


def parse_reference_numeric(ref_raw: str) -> Optional[ReferenceEntry]:
    r = (ref_raw or "").strip()
    if not r:
        return None

    m = re.match(r"^\s*\[\s*(\d+)\s*\]\s*(.+)$", r)
    if m:
        n = int(m.group(1))
        return ReferenceEntry(raw=r, key=key_numeric(n), pretty=f"[{n}]", number=n, author_or_org="", year="")

    m = re.match(r"^\s*(\d+)\s*[\.\)]\s*(.+)$", r)
    if m:
        n = int(m.group(1))
        return ReferenceEntry(raw=r, key=key_numeric(n), pretty=str(n), number=n, author_or_org="", year="")

    m = re.match(r"^\s*(\d+)\s+(.+)$", r)
    if m:
        n = int(m.group(1))
        return ReferenceEntry(raw=r, key=key_numeric(n), pretty=str(n), number=n, author_or_org="", year="")

    return None


def parse_reference_author_year(ref_raw: str) -> Optional[ReferenceEntry]:
    """
    Minimal parsing so cross-check works.
    Supports:
      Surname, ... (2020).
      Surname, ... 2020.
      WHO (2020) / World Health Organisation (2020) / UN (2020)
    """
    r = (ref_raw or "").strip()
    if not r:
        return None

    base = _strip_leading_numbering(r)

    # org first
    m_org = re.match(rf"^\s*([A-Z][A-Za-z&.\- ]{{2,}}?)\s*[\.\,]?\s*\(?\s*({YEAR})\s*\)?", base)
    if m_org:
        org = m_org.group(1).strip()
        y = m_org.group(2).strip()
        if is_known_org(org) and not looks_like_two_authors(org):
            k = key_org_year(org, y)
            return ReferenceEntry(raw=r, key=k, pretty=f"{titleish(org)} {y}", author_or_org=org, year=y)

    # author first
    m_au = re.match(rf"^\s*([A-Z][A-Za-z\-']+)\s*,.*?\b({YEAR})\b", base)
    if m_au:
        au = m_au.group(1).strip()
        y = m_au.group(2).strip()
        if looks_like_person_surname(au):
            k = key_author_year(au, y)
            return ReferenceEntry(raw=r, key=k, pretty=f"{titleish(au)} {y}", author_or_org=au, year=y)

    return None


# =============================
# In-text citation extraction (APA/Harvard)
# =============================
def extract_author_year_citations(text: str) -> List[InTextCitation]:
    out: List[InTextCitation] = []

    narr_people = re.compile(
        rf"""
        \b
        (?P<authors>
            [A-Z][A-Za-z\-']+
            (?:\s*,\s*[A-Z][A-Za-z\-']+)*
            (?:\s*,?\s*(?:and|&)\s*[A-Z][A-Za-z\-']+)?
            |
            [A-Z][A-Za-z\-']+\s+et\s+al\.
        )
        \s*
        \(\s*(?P<year>{YEAR})\s*\)
        """,
        flags=re.VERBOSE,
    )

    for m in narr_people.finditer(text or ""):
        authors_blob = m.group("authors").strip()
        y = m.group("year")

        authors_blob_clean = re.sub(r"\s+et\s+al\.\s*$", "", authors_blob, flags=re.IGNORECASE).strip()
        first = _first_surname_from_author_blob(authors_blob_clean) or authors_blob_clean.split(",")[0].strip()

        if not looks_like_person_surname(first):
            continue

        disp_authors = format_author_blob_for_display(authors_blob)
        key = key_author_year(first, y)
        pretty = f"{disp_authors} ({y})"
        out.append(InTextCitation("author-year", m.group(0), key, pretty, first, y))

    narr_org = re.compile(rf"\b([A-Z][A-Za-z&.\- ]{{2,}}?)\s*\(\s*({YEAR})\s*\)")
    for m in narr_org.finditer(text or ""):
        org, y = m.group(1).strip(), m.group(2)
        if looks_like_two_authors(org):
            continue
        if is_known_org(org):
            key = key_org_year(org, y)
            pretty = f"{titleish(org)} ({y})"
            out.append(InTextCitation("org-year", m.group(0), key, pretty, org, y))

    paren_block = re.compile(rf"\(([^()]*?\b{YEAR}\b[^()]*)\)")
    for m in paren_block.finditer(text or ""):
        block = m.group(1)
        parts = [p.strip() for p in block.split(";") if p.strip()]
        for p in parts:
            yrm = re.search(rf"\b({YEAR})\b", p)
            if not yrm:
                continue
            y = yrm.group(1)

            left = p
            m_left = re.search(rf"^(.+?)\s*,\s*{YEAR}\b", p)
            if m_left:
                left = m_left.group(1).strip()

            if is_known_org(left) and not looks_like_two_authors(left):
                key = key_org_year(left, y)
                out.append(InTextCitation("org-year", f"({p})", key, f"{titleish(left)} ({y})", left, y))
                continue

            first_author = _first_surname_from_author_blob(left) or _first_surname_from_author_blob(p)
            if not first_author or not looks_like_person_surname(first_author):
                continue

            key = key_author_year(first_author, y)
            pretty = f"{titleish(first_author)} ({y})"
            out.append(InTextCitation("author-year", f"({p})", key, pretty, first_author, y))

    return out


# =============================
# Vancouver / IEEE numeric extraction
# =============================
_SUP_DIGITS = {
    "⁰": "0", "¹": "1", "²": "2", "³": "3", "⁴": "4",
    "⁵": "5", "⁶": "6", "⁷": "7", "⁸": "8", "⁹": "9",
}


def _sup_to_int(s: str) -> Optional[int]:
    try:
        return int("".join(_SUP_DIGITS.get(ch, "") for ch in s))
    except Exception:
        return None


def _expand_numeric_chunks(inside: str) -> List[int]:
    chunks = [c.strip() for c in (inside or "").split(",") if c.strip()]
    nums: List[int] = []
    for c in chunks:
        r = re.match(r"^(\d+)\s*[-–]\s*(\d+)$", c)
        if r:
            a, b = int(r.group(1)), int(r.group(2))
            if a <= b and (b - a) <= 2000:
                nums.extend(range(a, b + 1))
        else:
            if c.isdigit():
                nums.append(int(c))
    return nums


def extract_ieee_numeric_citations(text: str) -> List[InTextCitation]:
    out: List[InTextCitation] = []
    open_b = r"[\[\［]"
    close_b = r"[\]\］]"

    pat = re.compile(
        rf"{open_b}"
        rf"(\s*\d+(?:\s*[-–]\s*\d+)?(?:\s*,\s*\d+(?:\s*[-–]\s*\d+)?)*)"
        rf"\s*{close_b}"
        rf"(?=[\s\.,;:\)\]\}}!?]|$)"
    )

    for m in pat.finditer(text or ""):
        raw = m.group(0)
        inside = m.group(1)
        for n in _expand_numeric_chunks(inside):
            out.append(InTextCitation("numeric", raw, key_numeric(n), f"[{n}]", number=n))

    return out


def extract_vancouver_numeric_citations(
    text: str,
    allow_plain_clusters: bool = True,
    strict_plain_clusters: bool = False,
) -> List[InTextCitation]:
    out: List[InTextCitation] = []

    # ---- (1) style ----
    open_p = r"[\(\（]"
    close_p = r"[\)\）]"
    paren_pat = re.compile(
        rf"{open_p}"
        rf"(\s*\d+(?:\s*[-–]\s*\d+)?(?:\s*,\s*\d+(?:\s*[-–]\s*\d+)?)*)"
        rf"\s*{close_p}"
        rf"(?=[\s\.,;:\)\]\}}!?]|$)"
    )

    for m in paren_pat.finditer(text or ""):
        raw = m.group(0)
        inside = m.group(1)
        if re.fullmatch(rf"\s*{YEAR}\s*", inside):
            continue
        for n in _expand_numeric_chunks(inside):
            out.append(InTextCitation("numeric", raw, key_numeric(n), f"({n})", number=n))

    # ---- [1] style ----
    out.extend(extract_ieee_numeric_citations(text))

    # ---- superscript runs ----
    sup_run = re.compile(r"[⁰¹²³⁴⁵⁶⁷⁸⁹]+")
    for m in sup_run.finditer(text or ""):
        s = m.group(0)
        next_ch = (text[m.end()] if m.end() < len(text) else "")
        if next_ch and not (next_ch.isspace() or next_ch in ".,;:)]}!?"):
            continue
        n = _sup_to_int(s)
        if n is None:
            continue
        out.append(InTextCitation("numeric", s, key_numeric(n), f"{n}", number=n))

    # superscript ranges like ¹–³
    sup_range = re.compile(r"([⁰¹²³⁴⁵⁶⁷⁸⁹]+)\s*[-–]\s*([⁰¹²³⁴⁵⁶⁷⁸⁹]+)")
    for m in sup_range.finditer(text or ""):
        a = _sup_to_int(m.group(1))
        b = _sup_to_int(m.group(2))
        if a is None or b is None:
            continue
        if a <= b and (b - a) <= 2000:
            raw = m.group(0)
            for n in range(a, b + 1):
                out.append(InTextCitation("numeric", raw, key_numeric(n), f"{n}", number=n))

    # plain digit clusters (lost superscript formatting)
    if allow_plain_clusters:
        plain_cluster_pat = re.compile(
            r"""
            (?x)
            (?P<prefix>[A-Za-z\)\]\}”"'])
            (?P<cluster>
                \d{1,4}
                (?:\s*[-–]\s*\d{1,4}
                 |\s*,\s*\d{1,4}
                 |\s+\d{1,4}
                )+
                |
                \d{1,4}\s*[-–]\s*\d{1,4}
            )
            (?=[\s\.,;:\)\]\}!?]|$)
            """,
            re.VERBOSE,
        )

        def _parse_plain_cluster(cluster: str) -> List[int]:
            c = (cluster or "").strip().replace("–", "-")
            c = re.sub(r"\s+", " ", c).strip()

            if strict_plain_clusters and re.search(r"\d\s+\d", c) and "," not in c and "-" not in c:
                return []

            if " " in c and "," not in c and "-" not in c:
                c = ",".join([x for x in c.split(" ") if x.strip()])

            c = c.replace(" ,", ",").replace(", ", ",")

            if "," in c:
                return _expand_numeric_chunks(c)

            m2 = re.match(r"^(\d+)\s*-\s*(\d+)$", c)
            if m2:
                a, b = int(m2.group(1)), int(m2.group(2))
                if a <= b and (b - a) <= 2000:
                    return list(range(a, b + 1))

            if c.isdigit():
                return [int(c)]
            return []

        for m in plain_cluster_pat.finditer(text or ""):
            raw_cluster = m.group("cluster")
            if re.search(r"\b(19|20)\d{2}\b", raw_cluster):
                continue
            if re.search(r"\d\.\d", raw_cluster):
                continue
            nums = _parse_plain_cluster(raw_cluster)
            for n in nums:
                if 1 <= n <= 5000:
                    out.append(InTextCitation("numeric", raw_cluster, key_numeric(n), f"{n}", number=n))

    return out


# =============================
# Year mismatch check
# =============================
def base_key(k: str) -> str:
    parts = k.split("_")
    return "_".join(parts[:-1]) if len(parts) >= 3 else k


def year_from_key(k: str) -> Optional[str]:
    parts = k.split("_")
    return parts[-1] if len(parts) >= 3 else None


def find_year_mismatches(cite_keys: List[str], ref_keys: List[str]) -> pd.DataFrame:
    cite_map = defaultdict(set)
    ref_map = defaultdict(set)

    for ck in cite_keys:
        b, y = base_key(ck), year_from_key(ck)
        if y:
            cite_map[b].add(y)

    for rk in ref_keys:
        b, y = base_key(rk), year_from_key(rk)
        if y:
            ref_map[b].add(y)

    rows = []
    for b in sorted(set(cite_map.keys()) & set(ref_map.keys())):
        if cite_map[b] != ref_map[b]:
            rows.append(
                {
                    "author_or_org_key": b,
                    "years_in_text": ", ".join(sorted(cite_map[b])),
                    "years_in_references": ", ".join(sorted(ref_map[b])),
                }
            )
    return pd.DataFrame(rows)


# =============================
# Online verification helpers
# =============================
def build_session() -> requests.Session:
    s = requests.Session()
    s.headers.update({"User-Agent": "citation-crosscheck/1.8"})
    return s


SESSION = build_session()


@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=1, max=7))
def _get_json(url: str, params: dict) -> dict:
    r = SESSION.get(url, params=params, timeout=25)
    r.raise_for_status()
    return r.json()


def extract_doi(text: str) -> Optional[str]:
    if not text:
        return None

    m = re.search(r"https?://doi\.org/(10\.\d{4,9}/[^\s<>\"]+)", text, flags=re.IGNORECASE)
    if m:
        doi = m.group(1)
    else:
        m2 = re.search(r"(10\.\d{4,9}/[^\s<>\"]+)", text, flags=re.IGNORECASE)
        if not m2:
            return None
        doi = m2.group(1)

    doi = doi.strip().strip(").,;:]}>\"'")
    doi = re.sub(r"&quot;|&gt;|&lt;|&amp;", "", doi)
    doi = doi.strip().strip(").,;:]}>\"'")
    return doi if doi.lower().startswith("10.") else None


@st.cache_data(show_spinner=False)
def crossref_lookup_by_doi(doi: str) -> Optional[dict]:
    try:
        data = _get_json(f"{CROSSREF_API}/{doi}", params={})
        return data.get("message")
    except Exception:
        return None


@st.cache_data(show_spinner=False)
def crossref_search(query: str, rows: int = 10) -> List[dict]:
    data = _get_json(CROSSREF_API, params={"query.bibliographic": query, "rows": rows})
    return data.get("message", {}).get("items", []) or []


@st.cache_data(show_spinner=False)
def openalex_search(query: str, per_page: int = 10) -> List[dict]:
    data = _get_json(OPENALEX_API, params={"search": query, "per-page": per_page})
    return data.get("results", []) or []


def _as_int_year(y: Optional[str]) -> Optional[int]:
    if not y:
        return None
    m = re.search(r"(19|20)\d{2}", str(y))
    return int(m.group(0)) if m else None


def crossref_item_year(it: dict) -> Optional[int]:
    issued = it.get("issued", {}).get("date-parts", [])
    if issued and issued[0]:
        try:
            return int(issued[0][0])
        except Exception:
            pass
    published_print = it.get("published-print", {}).get("date-parts", [])
    if published_print and published_print[0]:
        try:
            return int(published_print[0][0])
        except Exception:
            pass
    created = it.get("created", {}).get("date-parts", [])
    if created and created[0]:
        try:
            return int(created[0][0])
        except Exception:
            pass
    return None


def crossref_first_author_family(it: dict) -> str:
    authors = it.get("author") or []
    fam = authors[0].get("family") if authors else ""
    return fam or ""


def crossref_title(it: dict) -> str:
    t = it.get("title") or []
    return (t[0] if t else "") or ""


def openalex_year(it: dict) -> Optional[int]:
    y = it.get("publication_year")
    try:
        return int(y) if y else None
    except Exception:
        return None


def openalex_first_author_family(it: dict) -> str:
    authorships = it.get("authorships") or []
    if not authorships:
        return ""
    dn = (authorships[0].get("author") or {}).get("display_name") or ""
    return (dn.split()[-1] if dn else "") or ""


def openalex_title(it: dict) -> str:
    return (it.get("title") or "") or ""


def openalex_doi(it: dict) -> str:
    ids = it.get("ids") or {}
    d = ids.get("doi") or ""
    if d:
        d = d.replace("https://doi.org/", "").strip()
        d = d.strip(").,;:]}>\"'")
    return d


def author_match_ok(ref_surname: str, cand_surname: str) -> bool:
    if not ref_surname or not cand_surname:
        return False
    return norm_token(ref_surname) == norm_token(cand_surname)


def year_match_ok(ref_year: Optional[int], cand_year: Optional[int]) -> bool:
    if ref_year is None or cand_year is None:
        return False
    if ref_year == cand_year:
        return True
    return abs(ref_year - cand_year) == 1


def guess_title_snippet(ref: str) -> str:
    r = _strip_leading_numbering(ref)
    t = re.sub(rf"\(.*?{YEAR}.*?\)", " ", r)
    t = re.sub(r"^[^\.]{1,220}\.\s*", " ", t)
    t = norm_spaces(t)
    words = t.split()
    return " ".join(words[:22])[:260]


def verify_reference_online(
    ref_obj: ReferenceEntry,
    throttle_s: float = 0.2,
    use_crossref: bool = True,
    use_openalex: bool = True,
) -> dict:
    ref_entry = ref_obj.raw or ""
    doi = extract_doi(ref_entry)

    ref_year = _as_int_year(ref_obj.year)
    who = (ref_obj.author_or_org or "").strip()
    who_is_org = ref_obj.key.startswith("org_")
    ref_surname = who if (who and (not who_is_org) and looks_like_person_surname(who)) else ""

    title_snip = guess_title_snippet(ref_entry)

    if doi and use_crossref:
        cr = crossref_lookup_by_doi(doi)
        if cr:
            y = crossref_item_year(cr)
            fam = crossref_first_author_family(cr)
            title = crossref_title(cr)
            return {
                "status": "verified",
                "source": "crossref_doi",
                "score": 100,
                "doi": doi,
                "matched_year": str(y or ""),
                "matched_first_author": fam,
                "matched_title": (title or "")[:180],
                "query_used": "doi_lookup",
                "error_crossref": "",
                "error_openalex": "",
            }
        return {
            "status": "not_found",
            "source": "crossref_doi",
            "score": 0,
            "doi": doi,
            "matched_year": "",
            "matched_first_author": "",
            "matched_title": "",
            "query_used": "doi_lookup",
            "error_crossref": "",
            "error_openalex": "",
        }

    parts = []
    if who:
        parts.append(who)
    if ref_year:
        parts.append(str(ref_year))
    if title_snip:
        parts.append(title_snip)

    query = " ".join(parts).strip() or _strip_leading_numbering(ref_entry)[:200]
    time.sleep(max(0.0, float(throttle_s or 0.0)))

    best = {
        "status": "not_found",
        "source": "",
        "score": 0,
        "doi": "",
        "matched_year": "",
        "matched_first_author": "",
        "matched_title": "",
        "query_used": query[:220],
        "error_crossref": "",
        "error_openalex": "",
    }

    if use_crossref:
        try:
            items = crossref_search(query, rows=10)
            for it in items:
                cand_year = crossref_item_year(it)
                cand_fam = crossref_first_author_family(it)
                cand_title = crossref_title(it)
                cand_doi = (it.get("DOI") or "").strip()

                if not who_is_org:
                    if ref_surname and not author_match_ok(ref_surname, cand_fam):
                        continue
                    if ref_year and not year_match_ok(ref_year, cand_year):
                        continue
                else:
                    if ref_year and not year_match_ok(ref_year, cand_year):
                        continue

                score = fuzz.WRatio(title_snip, cand_title) if (title_snip and cand_title) else 70
                status = "verified" if score >= 90 else ("likely" if score >= 82 else "needs_review")

                if score > best["score"]:
                    best = {
                        "status": status,
                        "source": "crossref",
                        "score": int(score),
                        "doi": cand_doi,
                        "matched_year": str(cand_year or ""),
                        "matched_first_author": cand_fam,
                        "matched_title": (cand_title or "")[:180],
                        "query_used": query[:220],
                        "error_crossref": "",
                        "error_openalex": best.get("error_openalex", ""),
                    }
        except Exception as e:
            best["error_crossref"] = str(e)[:220]

    if use_openalex:
        try:
            items = openalex_search(query, per_page=10)
            for it in items:
                cand_year = openalex_year(it)
                cand_fam = openalex_first_author_family(it)
                cand_title = openalex_title(it)
                cand_doi = openalex_doi(it)

                if not who_is_org:
                    if ref_surname and not author_match_ok(ref_surname, cand_fam):
                        continue
                    if ref_year and not year_match_ok(ref_year, cand_year):
                        continue
                else:
                    if ref_year and not year_match_ok(ref_year, cand_year):
                        continue

                score = fuzz.WRatio(title_snip, cand_title) if (title_snip and cand_title) else 70
                status = "verified" if score >= 90 else ("likely" if score >= 82 else "needs_review")

                if score > best["score"]:
                    best = {
                        "status": status,
                        "source": "openalex",
                        "score": int(score),
                        "doi": cand_doi,
                        "matched_year": str(cand_year or ""),
                        "matched_first_author": cand_fam,
                        "matched_title": (cand_title or "")[:180],
                        "query_used": query[:220],
                        "error_crossref": best.get("error_crossref", ""),
                        "error_openalex": "",
                    }
        except Exception as e:
            best["error_openalex"] = str(e)[:220]

    return best


# =============================
# Reference-style detection + reformat + DOCX export
# =============================
def detect_reference_style(ref_entries: List[str]) -> Tuple[str, Dict[str, int]]:
    counts = {STYLE_APA: 0, STYLE_HARVARD: 0, STYLE_NUMERIC: 0}

    for raw in ref_entries or []:
        r = (raw or "").strip()
        if not r:
            continue

        if re.match(r"^\s*(\[\d+\]|\d+[\.\)])\s+", r) or re.match(r"^\s*\d+\s+[A-Z]", r):
            counts[STYLE_NUMERIC] += 1
            continue

        if re.search(rf"^[A-Z][A-Za-z\-']+\s*,.+?\(\s*{YEAR}\s*\)\.", r):
            counts[STYLE_APA] += 1
            continue

        if re.search(rf"^[A-Z][A-Za-z\-']+\s*,.+?\b{YEAR}\b", r) and "(" not in r[:60]:
            counts[STYLE_HARVARD] += 1
            continue

    dominant = max(counts.items(), key=lambda kv: kv[1])[0]
    return dominant, counts


def _get_first_author_and_year_from_ref(ref_raw: str) -> Tuple[str, str]:
    e = _strip_leading_numbering(ref_raw)

    m_org = re.match(r"^\s*([A-Z][A-Za-z&.\- ]{2,}?)\s*[\.\,]?\s*\(?\s*(" + YEAR + r")\s*\)?", e)
    if m_org:
        who = m_org.group(1).strip()
        y = m_org.group(2).strip()
        return titleish(who), y

    m_au = re.match(r"^\s*([A-Z][A-Za-z\-']+)\s*,.*?\(?\s*(" + YEAR + r")\s*\)?", e)
    if m_au:
        return titleish(m_au.group(1).strip()), m_au.group(2).strip()

    y = ""
    ym = re.search(rf"\b({YEAR})\b", e)
    if ym:
        y = ym.group(1)
    return "Unknown", y


def format_reference_entry(ref_raw: str, target_style: str, idx: int = 1, enrich: bool = False) -> str:
    base = _strip_leading_numbering(ref_raw).strip()

    if target_style == STYLE_APA:
        if re.search(rf"\(\s*{YEAR}\s*\)", base):
            return base
        who, y = _get_first_author_and_year_from_ref(base)
        if y:
            return re.sub(rf"\b{re.escape(y)}\b", f"({y}).", base, count=1)
        return base

    if target_style == STYLE_HARVARD:
        if re.search(rf"\(\s*{YEAR}\s*\)", base):
            return re.sub(rf"\(\s*({YEAR})\s*\)", r"\1", base)
        return base

    if target_style == STYLE_NUMERIC:
        if (
            re.match(r"^\s*\[\d+\]", ref_raw)
            or re.match(r"^\s*\d+[\.\)]", ref_raw)
            or re.match(r"^\s*\d+\s+", ref_raw)
        ):
            return ref_raw
        return f"[{idx}] {base}"

    return base


def export_references_to_docx(formatted_refs: List[str], title: str = "References") -> bytes:
    if DocxDocument is None:
        raise RuntimeError("python-docx not available for DOCX export")
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    for r in formatted_refs:
        doc.add_paragraph(r)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =============================
# PDF report
# =============================
def make_pdf_report(
    style_name: str,
    summary: dict,
    df_missing: pd.DataFrame,
    df_uncited: pd.DataFrame,
    df_mismatch: pd.DataFrame,
    df_verify: pd.DataFrame,
) -> bytes:
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 2 * cm

    c.setFont("Helvetica-Bold", 14)
    c.drawString(2 * cm, y, "Citation Crosscheck Report")
    y -= 0.8 * cm

    c.setFont("Helvetica", 10)
    c.drawString(2 * cm, y, f"Style: {style_name}")
    y -= 0.8 * cm

    c.setFont("Helvetica-Bold", 11)
    c.drawString(2 * cm, y, "Summary")
    y -= 0.6 * cm

    c.setFont("Helvetica", 10)
    for k, v in summary.items():
        c.drawString(2 * cm, y, f"{k.replace('_',' ').title()}: {v}")
        y -= 0.5 * cm
        if y < 2 * cm:
            c.showPage()
            y = height - 2 * cm

    def draw_table(title: str, df: pd.DataFrame, max_rows: int = 30):
        nonlocal y
        c.showPage()
        y = height - 2 * cm
        c.setFont("Helvetica-Bold", 12)
        c.drawString(2 * cm, y, title)
        y -= 0.7 * cm
        c.setFont("Helvetica", 9)

        if df is None or df.empty:
            c.drawString(2 * cm, y, "None detected.")
            return

        cols = list(df.columns)
        c.drawString(2 * cm, y, " | ".join(cols))
        y -= 0.4 * cm
        c.line(2 * cm, y, width - 2 * cm, y)
        y -= 0.3 * cm

        for _, row in df.head(max_rows).iterrows():
            line = " | ".join(str(row[col])[:72] for col in cols)
            c.drawString(2 * cm, y, line)
            y -= 0.45 * cm
            if y < 2 * cm:
                c.showPage()
                y = height - 2 * cm
                c.setFont("Helvetica", 9)

        if len(df) > max_rows:
            c.drawString(2 * cm, y, f"... showing {max_rows} of {len(df)} rows")
            y -= 0.6 * cm

    draw_table("Cited in-text but missing in References", df_missing)
    draw_table("In References but never cited", df_uncited)
    draw_table("Year Mismatches", df_mismatch)
    draw_table("Online Verification", df_verify)

    c.save()
    return buf.getvalue()


# =============================
# Streamlit UI
# =============================
st.set_page_config(page_title="Citation Crosschecker", layout="wide")
st.title("Citation Crosschecker")

st.markdown(
    """
<div style="margin-top:6px; margin-bottom:10px; padding:10px 12px; border:1px solid #e6e6e6; border-radius:10px;">
  <div style="font-size:13px; font-weight:600;">
    © Prof. Anokye M. Adam
  </div>
  <div style="font-size:12px; margin-top:6px; line-height:1.35;">
    Disclaimer: This tool can make mistakes, including missed or incorrect citation matches and online metadata errors.
    Always verify results against your manuscript, style guide, and original sources before submission.
  </div>
</div>
""",
    unsafe_allow_html=True,
)

col1, col2 = st.columns(2)
with col1:
    uploaded = st.file_uploader("Upload DOCX, PDF, or TXT", type=["docx", "pdf", "txt"])
with col2:
    pasted = st.text_area("Or paste your manuscript text", height=170)

style = st.selectbox(
    "Citation style to check",
    ["APA/Harvard (author–year)", "IEEE (numeric [1])", "Vancouver (numeric 1)"],
)

# Vancouver options
v_allow_plain = True
v_strict_plain = False
if style.startswith("Vancouver"):
    st.caption("Vancouver in-text can be (1), [1], superscripts, or plain digits from PDF extraction.")
    v_allow_plain = st.checkbox(
        "Capture plain numeric clusters like 1,2,3 or 1-3 (useful when superscripts are lost)",
        value=True,
    )
    v_strict_plain = st.checkbox(
        "Strict mode (ignore space-separated lists like 1 2 3)",
        value=False,
        disabled=not v_allow_plain,
    )

text = ""
if uploaded is not None:
    name = uploaded.name.lower()
    try:
        if name.endswith(".docx"):
            text = read_docx(uploaded)
        elif name.endswith(".pdf"):
            text = read_pdf(uploaded)
        elif name.endswith(".txt"):
            text = uploaded.read().decode("utf-8", errors="ignore")
    except Exception as e:
        st.error(f"Could not read file: {e}")
elif pasted.strip():
    text = pasted

if not text.strip():
    st.info("Upload a file or paste text to begin.")
    st.stop()

increment_once_per_session("counted_app_run", app_run=True)

main_text, ref_text, ref_msg = split_by_heading_or_autodetect(text)

st.subheader("References detection")
st.write(ref_msg)

auto_idx, auto_conf = auto_detect_references_start(text)

force_manual = ("No heading found" in ref_msg) and (auto_conf < 0.60)
manual = st.checkbox(
    "Manually choose where References start (recommended if heading not found)",
    value=force_manual or ("No heading found" in ref_msg),
)

if manual:
    lines = text.splitlines()
    guess = auto_idx if auto_idx is not None else max(0, int(len(lines) * 0.70))
    idx = st.slider("Select the line where the References section starts", 0, max(0, len(lines) - 1), guess)
    main_text = "\n".join(lines[:idx]).strip()
    ref_text = "\n".join(lines[idx:]).strip()

st.caption(f"Main text length: {len(main_text):,} chars | References length: {len(ref_text):,} chars")

with st.expander("Preview detected References section"):
    st.text(ref_text[:8000] if ref_text else "No references detected yet")

if not ref_text.strip():
    st.warning("References section is empty. Paste your References or adjust the slider.")
    ref_text = st.text_area("Paste References section here", height=240)

# Extract citations and references
ref_raw = split_reference_entries(ref_text)

if style.startswith("APA/Harvard"):
    cites = extract_author_year_citations(main_text)
    refs = [parse_reference_author_year(r) for r in ref_raw]
    refs = [r for r in refs if r is not None]
elif style.startswith("IEEE"):
    cites = extract_ieee_numeric_citations(main_text)
    refs = [parse_reference_numeric(r) for r in ref_raw]
    refs = [r for r in refs if r is not None]
else:
    cites = extract_vancouver_numeric_citations(
        main_text,
        allow_plain_clusters=v_allow_plain,
        strict_plain_clusters=v_strict_plain,
    )
    refs = [parse_reference_numeric(r) for r in ref_raw]
    refs = [r for r in refs if r is not None]

increment_once_per_session(
    "counted_file_run",
    file=True,
    cites=len(cites),
    refs=len(refs),
)

# Reference style tools
st.divider()
st.subheader("Reference style check and reformat (export to Word)")

det_style, det_counts = detect_reference_style(ref_raw)
st.write(
    f"Detected dominant reference style: **{det_style}** "
    f"(APA-like: {det_counts[STYLE_APA]}, Harvard-like: {det_counts[STYLE_HARVARD]}, Numeric: {det_counts[STYLE_NUMERIC]})"
)

target_style = st.selectbox(
    "Reformat all reference entries into",
    [STYLE_APA, STYLE_HARVARD, STYLE_NUMERIC],
    index=[STYLE_APA, STYLE_HARVARD, STYLE_NUMERIC].index(det_style)
    if det_style in [STYLE_APA, STYLE_HARVARD, STYLE_NUMERIC]
    else 0,
)

enrich_with_doi = st.checkbox(
    "Improve formatting using DOI metadata (Crossref lookup when DOI exists)",
    value=False,
)

if st.button("Generate reformatted References list"):
    formatted = [
        format_reference_entry(raw, target_style, idx=i, enrich=enrich_with_doi)
        for i, raw in enumerate(ref_raw, start=1)
    ]

    st.markdown("#### Preview (first 20)")
    st.text("\n".join(formatted[:20]))

    if DocxDocument is None:
        st.warning("DOCX export needs python-docx installed in your environment.")
    else:
        docx_bytes = export_references_to_docx(formatted, title="References")
        st.download_button(
            "Download reformatted References (DOCX)",
            data=docx_bytes,
            file_name=f"references_reformatted_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    st.download_button(
        "Download reformatted References (TXT)",
        data="\n".join(formatted).encode("utf-8"),
        file_name="references_reformatted.txt",
        mime="text/plain",
    )

# =============================
# Crosscheck tables (FIXED grouping + better display)
# =============================
cite_keys = [c.key for c in cites]
ref_keys = [r.key for r in refs]

missing_keys = sorted(set(cite_keys) - set(ref_keys))
uncited_keys = sorted(set(ref_keys) - set(cite_keys))

cite_counts = Counter(cite_keys)
ref_counts = Counter(ref_keys)

# choose stable display text per key
key_to_cite_pretty = {}
for c in cites:
    if c.key not in key_to_cite_pretty:
        key_to_cite_pretty[c.key] = c.pretty or c.raw or c.key

key_to_full_ref = {}
for r in refs:
    if r.key not in key_to_full_ref:
        key_to_full_ref[r.key] = r.raw or r.pretty or r.key

suggestions: Dict[str, List[str]] = {}
if missing_keys and ref_keys:
    # suggest based on key similarity, but show full reference text
    for mk in missing_keys:
        matches = process.extract(mk, ref_keys, scorer=fuzz.WRatio, limit=5)
        best = []
        for cand_key, score, _ in matches:
            if score >= 75:
                best.append(f"{key_to_full_ref.get(cand_key, cand_key)[:160]} ({int(score)})")
        suggestions[mk] = best

df_missing = pd.DataFrame(
    [
        {
            "citation_key": k,
            "citation_in_text": key_to_cite_pretty.get(k, k),
            "count_in_text": cite_counts[k],
            "suggested_matches": " | ".join(suggestions.get(k, [])),
        }
        for k in missing_keys
    ]
)

df_uncited = pd.DataFrame(
    [
        {
            "reference_key": k,
            "reference_full": key_to_full_ref.get(k, k),
            "times_in_references": ref_counts[k],
        }
        for k in uncited_keys
    ]
)

df_mismatch = pd.DataFrame()
if style.startswith("APA/Harvard"):
    df_mismatch = find_year_mismatches(cite_keys, ref_keys)

summary = {
    "in_text_citations_found": len(cites),
    "reference_entries_found": len(refs),
    "missing_in_references": len(missing_keys),
    "uncited_references": len(uncited_keys),
}

m1, m2, m3, m4 = st.columns(4)
m1.metric("In-text citations", summary["in_text_citations_found"])
m2.metric("Reference entries", summary["reference_entries_found"])
m3.metric("Missing", summary["missing_in_references"])
m4.metric("Uncited", summary["uncited_references"])

st.divider()
c1, c2 = st.columns(2)

with c1:
    st.markdown("### Cited in-text but missing in References")
    if df_missing.empty:
        st.info("None detected.")
    else:
        st.dataframe(df_missing, use_container_width=True)
    st.download_button(
        "Download missing (CSV)",
        df_missing.to_csv(index=False).encode("utf-8"),
        file_name="missing_in_references.csv",
        mime="text/csv",
    )

with c2:
    st.markdown("### In References but never cited")
    if df_uncited.empty:
        st.info("None detected.")
    else:
        st.dataframe(df_uncited, use_container_width=True)
    st.download_button(
        "Download uncited (CSV)",
        df_uncited.to_csv(index=False).encode("utf-8"),
        file_name="uncited_references.csv",
        mime="text/csv",
    )

st.divider()
st.markdown("### Year mismatches (same author/org, different year)")
if style.startswith("APA/Harvard"):
    if df_mismatch.empty:
        st.success("No year mismatches detected.")
    else:
        st.dataframe(df_mismatch, use_container_width=True)
        st.download_button(
            "Download mismatches (CSV)",
            df_mismatch.to_csv(index=False).encode("utf-8"),
            file_name="year_mismatches.csv",
            mime="text/csv",
        )
else:
    st.info("Year mismatch check applies to author–year styles only.")

# =============================
# Online verification
# =============================
st.divider()
st.subheader("Online verification (Crossref + OpenAlex)")

enable_verify = st.checkbox("Enable online verification", value=False)
use_crossref = st.checkbox("Use Crossref", value=True, disabled=not enable_verify)
use_openalex = st.checkbox("Use OpenAlex", value=True, disabled=not enable_verify)
throttle = st.slider("Throttle seconds between queries", 0.0, 2.0, 0.25, 0.05, disabled=not enable_verify)

df_verify = pd.DataFrame()

if enable_verify:
    test_ok = True
    test_msg = ""
    if use_crossref:
        try:
            _ = _get_json("https://api.crossref.org/works", {"rows": 1})
        except Exception as e:
            test_ok = False
            test_msg = str(e)[:220]

    if (use_crossref and not test_ok) and (not use_openalex):
        st.error("Online verification can’t reach Crossref from this deployment.")
        st.write(f"Error: {test_msg}")
    else:
        rows = []
        progress = st.progress(0)
        total = len(refs) if refs else 1

        for i, r in enumerate(refs):
            res = verify_reference_online(
                r,
                throttle_s=throttle,
                use_crossref=use_crossref,
                use_openalex=use_openalex,
            )
            rows.append(
                {
                    "reference": (r.raw[:220] + "…") if len(r.raw) > 220 else r.raw,
                    "status": res.get("status", ""),
                    "source": res.get("source", ""),
                    "score": res.get("score", ""),
                    "doi": res.get("doi", ""),
                    "matched_year": res.get("matched_year", ""),
                    "matched_first_author": res.get("matched_first_author", ""),
                    "matched_title": res.get("matched_title", ""),
                    "query_used": res.get("query_used", ""),
                    "error_crossref": res.get("error_crossref", ""),
                    "error_openalex": res.get("error_openalex", ""),
                }
            )
            progress.progress(int((i + 1) / total * 100))

        df_verify = pd.DataFrame(rows)
        st.dataframe(df_verify, use_container_width=True)
        st.download_button(
            "Download verification (CSV)",
            df_verify.to_csv(index=False).encode("utf-8"),
            file_name="online_verification.csv",
            mime="text/csv",
        )

        flagged = df_verify[df_verify["status"].isin(["not_found", "needs_review"])]
        if len(flagged) > 0:
            st.warning(f"Flagged {len(flagged)} items as Not found or Needs review. Check these first.")

# =============================
# PDF report
# =============================
st.divider()
pdf_bytes = make_pdf_report(style, summary, df_missing, df_uncited, df_mismatch, df_verify)
st.download_button(
    "Download full PDF report",
    data=pdf_bytes,
    file_name="citation_crosscheck_report.pdf",
    mime="application/pdf",
)

# =============================
# Usage statistics display
# =============================
st.divider()
st.subheader("Usage statistics")
stats = load_usage_stats()
u1, u2, u3, u4 = st.columns(4)
u1.metric("App runs", stats["app_runs"])
u2.metric("Files processed", stats["files_processed"])
u3.metric("Citations checked", stats["citations_checked"])
u4.metric("References checked", stats["references_checked"])

# =============================
# Debug
# =============================
with st.expander("Extracted items (debug)"):
    tab1, tab2, tab3 = st.tabs(["In-text citations", "Reference entries", "Raw reference splits"])
    with tab1:
        st.dataframe(pd.DataFrame([c.__dict__ for c in cites]), use_container_width=True)
    with tab2:
        st.dataframe(pd.DataFrame([r.__dict__ for r in refs]), use_container_width=True)
    with tab3:
        st.write(f"Split into {len(ref_raw)} raw entries")
        st.text("\n\n---\n\n".join(ref_raw[:20]))
